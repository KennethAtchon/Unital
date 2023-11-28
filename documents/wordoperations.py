import tkinter as tk
from tkinter import filedialog
import customtkinter
import os
from difflib import Differ
from docx import Document
import threading
import pyttsx3

class TextOperationsFrame(customtkinter.CTkTabview):
    def __init__(self, master):
        super().__init__(master, width=500)
        self.add("TextOp")
        self.add("Compare Files")
        self.add("Tab 3")
        self.tab("TextOp").grid_columnconfigure(0, weight=1)  # configure grid of individual tabs
        self.tab("Compare Files").grid_columnconfigure(0, weight=1)

        # textOp section
        self.op_label = customtkinter.CTkLabel(self.tab("TextOp"), text="Text Operations", font=customtkinter.CTkFont(size=14, weight="bold"))
        self.op_label.grid(row=0, column=0, padx=20, pady=20)

        self.op_textbox = customtkinter.CTkTextbox(self.tab("TextOp"), height=60)
        self.op_textbox.grid(row=1, column=0, padx=20, pady=(0, 20), sticky="nsew")

        self.select_file_button = customtkinter.CTkButton(self.tab("TextOp"), text="Select File", command=self.select_file)
        self.select_file_button.grid(row=1, column=1, padx=20, pady=(0, 20))

        self.checkboxes = []
        self.checkbox_vars = [tk.IntVar() for _ in range(3)]

        for i, (text, var) in enumerate(zip(["Word Count", "Reverse Text", "Text to Speech"], self.checkbox_vars)):
            cb = customtkinter.CTkCheckBox(self.tab("TextOp"), text=text, variable=var)
            cb.grid(row=i+2, column=0, padx=20, pady=(0, 10), sticky="w")
            self.checkboxes.append(cb)

        self.op_submit = customtkinter.CTkButton(self.tab("TextOp"), text="Process Text", command=self.handle_text_ops)
        self.op_submit.grid(row=5, column=0)

        self.file_path = None

        # Compare Files section
        self.com_label = customtkinter.CTkLabel(self.tab("Compare Files"), text="Select two files to compare", font=customtkinter.CTkFont(size=14, weight="bold"))
        self.com_label.grid(row=0, column=0, padx=20)

        self.file1_entry = customtkinter.CTkEntry(self.tab("Compare Files"), placeholder_text="File1")
        self.file1_entry.grid(row=1, column=0)
        self.file2_entry = customtkinter.CTkEntry(self.tab("Compare Files"), placeholder_text="File2")
        self.file2_entry.grid(row=2, column=0)
        self.file1_button = customtkinter.CTkButton(self.tab("Compare Files"), text="Browse", command=self.browse_file_1)
        self.file1_button.grid(row=1, column=1, padx=20, pady=(10,0))
        self.file2_button = customtkinter.CTkButton(self.tab("Compare Files"), text="Browse", command=self.browse_file_2)
        self.file2_button.grid(row=2, column=1, padx=20, pady=(10,0))
        self.compare_button = customtkinter.CTkButton(self.tab("Compare Files"), text="Compare Files", command=self.compare_files)
        self.compare_button.grid(row=3, column=0, padx=20, pady=20)
        self.diff_label = customtkinter.CTkLabel(self.tab("Compare Files"), text="Differences:  ??%")
        self.diff_label.grid(row=3, column=1, padx=20)


    def select_file(self):
        file_types = [("Text files", ".txt .docx")]
        file_path = filedialog.askopenfilename(filetypes=file_types)

        if file_path:
            _, file1_extension = os.path.splitext(file_path)
            if file1_extension.lower() == '.docx':
                self.file_path = file_path
                doc1 = Document(file_path)

                # Clear the existing content in the Text widget
                self.op_textbox.delete(1.0, "end")

                # Insert each paragraph's text into the Text widget
                for paragraph in doc1.paragraphs:
                    text = paragraph.text
                    self.op_textbox.insert("end", f"{text}\n")

            else:
                self.file_path = file_path
                with open(file_path, "r") as f:
                    text = f.read()
                    self.op_textbox.delete(1.0, "end")
                    self.op_textbox.insert(1.0, text)


    def handle_text_ops(self):
        if self.file_path:
            user_text = self.op_textbox.get("1.0", "end-1c")
        else:
            user_text = self.op_textbox.get("1.0", "end-1c") or self.read_file()
            self.file_path = ''

        output = []

        if self.checkbox_vars[0].get():
            word_count = len(user_text.split())
            character_count = len(user_text)
            #print(f"Word Count: {word_count}")
            output.append(f"Word Count: {word_count}")
            output.append(f"Character Count: {character_count}")

        if self.checkbox_vars[1].get():
            reversed_text = user_text[::-1]
            #print(f"Reversed: {reversed_text}")
            output.append(f"Reversed: {reversed_text}")

        if self.checkbox_vars[2].get():
            threading.Thread(target=self.perform_tts, args=(user_text,)).start()

        # Create a folder with the same name as the path
        folder_path = os.path.splitext(self.file_path)[0]
        os.makedirs(folder_path, exist_ok=True)

        # Create a text file with details
        text_details = "\n".join(output)
        text_file_path = os.path.join(folder_path, f'{self.file_path}details.txt')
        with open(text_file_path, 'w') as text_file:
            text_file.write(text_details)

        print(f"Text details saved to: {text_file_path}")

    def perform_tts(self, user_text):
        # Initialize the text-to-speech engine
        engine = pyttsx3.init()

        # Set properties (optional)
        engine.setProperty('rate', 150)  # Speed of speech

        # Save speech to a temporary WAV file
        wav_file_path = os.path.splitext(self.file_path)[0] + '.wav'
        engine.save_to_file(user_text, wav_file_path)

        # Wait for the speech to be generated
        engine.runAndWait()

        print(f"Speech saved to: {wav_file_path}")

    def browse_file_1(self):
        file_types = [("Accepted Files", ".txt .docx .py .java .html .css .js"), ("Other Files", "")]
        file_path = filedialog.askopenfilename(filetypes=file_types)
        self.file1_entry.delete(0, tk.END)
        self.file1_entry.insert(0, file_path)

    def browse_file_2(self):
        file_types = [("Accepted Files", ".txt .docx .py .java .html .css .js"),  ("Other Files", "")]
        file_path = filedialog.askopenfilename(filetypes=file_types)
        self.file2_entry.delete(0, tk.END)
        self.file2_entry.insert(0, file_path)



    def compare_files(self):
        file1_path = self.file1_entry.get()
        file2_path = self.file2_entry.get()

        try:
            # Check file extensions to decide how to read the files
            _, file1_extension = os.path.splitext(file1_path)
            _, file2_extension = os.path.splitext(file2_path)

            if file1_extension.lower() == '.docx' and file2_extension.lower() == '.docx':
                # Read content from .docx files
                doc1 = Document(file1_path)
                doc2 = Document(file2_path)

                # Extract text from paragraphs in the document
                lines1 = [f" {i + 1} {paragraph.text.rstrip()}" for i, paragraph in enumerate(doc1.paragraphs)]
                lines2 = [f" {i + 1} {paragraph.text.rstrip()}" for i, paragraph in enumerate(doc2.paragraphs)]
            else:
                # Read content from text files
                with open(file1_path, 'r', encoding='utf-8') as file1, open(file2_path, 'r', encoding='utf-8') as file2:
                    lines1 = [f" {i + 1} {line.rstrip()}" for i, line in enumerate(file1)]
                    lines2 = [f" {i + 1} {line.rstrip()}" for i, line in enumerate(file2)]

        except FileNotFoundError as e:
            print(f"Error: File not found: {e.filename}")
            return

        differ = Differ()
        diff = differ.compare(lines1, lines2)
        difference_count = 0
        result_lines = []

        for line in diff:
            if line.startswith("+"):
                difference_count += 1
                result_lines.append(f"{file1_path} + {line[1:]}")
            elif line.startswith("-"):
                difference_count += 1
                result_lines.append(f"{file2_path} - {line[1:]}")

        total_lines = len(lines1) + len(lines2)
        percentage_diff = (difference_count / total_lines) * 100
        result_lines.append(f"\nPercentage difference: {percentage_diff:.2f}%")
        self.diff_label.configure(text=f"Percentage Difference: {percentage_diff:.2f}%")

        # Create CompareResults folder
        folder_path = os.path.join(os.path.dirname(file1_path), 'CompareResults')
        os.makedirs(folder_path, exist_ok=True)

        # Write results to a text file in the CompareResults folder
        result_file_path = os.path.join(folder_path, 'CompareResults.txt')
        with open(result_file_path, 'w', encoding='utf-8') as result_file:
            result_file.write('\n'.join(result_lines))

        print(f"Comparison results saved to: {result_file_path}")

    def read_file(self):
        if self.file_path:
            with open(self.file_path, "r") as f:
                return f.read()
        return ""