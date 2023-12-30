import customtkinter
from tkinter import filedialog
import os
import subprocess
import tkinter as tk
import time
import docx
from PyPDF2 import PdfReader, PdfFileWriter

class ConvertFrame(customtkinter.CTkTabview):
    def __init__(self, master):
        super().__init__(master, width=250)
        self.add("Doc2Pdf")
        self.tab("Doc2Pdf").grid_columnconfigure(0, weight=1)
        
        self.d2p_label = customtkinter.CTkLabel(self.tab("Doc2Pdf"), text="Converts docx file to Pdf", font=customtkinter.CTkFont(size=14, weight="bold"))
        self.d2p_entry = customtkinter.CTkEntry(self.tab("Doc2Pdf"), placeholder_text="Enter the path to docx...")
        self.d2p_browse_button = customtkinter.CTkButton(self.tab("Doc2Pdf"), text="Browse", command=self.browse_file_d2p)
        self.d2p_browse_button.grid(row=2, column=1, padx=10, pady=(20, 0))
        self.d2psubmit_button = customtkinter.CTkButton(self.tab("Doc2Pdf"), text="Convert", command=self.convert_docx_to_pdf)
        self.d2p_label.grid(row=0, column=0)
        self.d2psubmit_button.grid(row=3, column=0, pady=(20,0))
        self.d2p_entry.grid(row=2, column=0, padx=(20, 0), pady=(20, 0),  sticky="nsew")

        self.progressbar_1 = customtkinter.CTkProgressBar(self.tab("Doc2Pdf"))
        self.progressbar_1.grid(row=4, column=0, pady=(20, 0), sticky="ew")
        self.progressbar_1.set(0)
        self.progressbar_1.configure(mode="determinate")


        # Add new tab for PDF to DOCX conversion
        self.add("Pdf2Doc")
        self.tab("Pdf2Doc").grid_columnconfigure(0, weight=1)
        
        self.p2d_label = customtkinter.CTkLabel(self.tab("Pdf2Doc"), text="Converts pdf file to Docx", font=customtkinter.CTkFont(size=14, weight="bold"))
        self.p2d_entry = customtkinter.CTkEntry(self.tab("Pdf2Doc"), placeholder_text="Enter the path to pdf...")
        self.p2d_browse_button = customtkinter.CTkButton(self.tab("Pdf2Doc"), text="Browse", command=self.browse_file_p2d)
        self.p2d_browse_button.grid(row=2, column=1, padx=10, pady=(20, 0))
        self.p2dsubmit_button = customtkinter.CTkButton(self.tab("Pdf2Doc"), text="Convert", command=self.convert_pdf_to_docx)
        self.p2d_label.grid(row=0, column=0)
        self.p2dsubmit_button.grid(row=3, column=0, pady=(20,0))
        self.p2d_entry.grid(row=2, column=0, padx=(20, 0), pady=(20, 0),  sticky="nsew")

        self.progressbar_2 = customtkinter.CTkProgressBar(self.tab("Pdf2Doc"))
        self.progressbar_2.grid(row=4, column=0, pady=(20, 0), sticky="ew")
        self.progressbar_2.set(0)
        self.progressbar_2.configure(mode="determinate")

        self.add("Doc2Txt")
        self.tab("Doc2Txt").grid_columnconfigure(0, weight=1)
        
        self.d2t_label = customtkinter.CTkLabel(self.tab("Doc2Txt"), text="Converts docx file to Txt", font=customtkinter.CTkFont(size=14, weight="bold"))
        self.d2t_entry = customtkinter.CTkEntry(self.tab("Doc2Txt"), placeholder_text="Enter the path to docx...")
        self.d2t_browse_button = customtkinter.CTkButton(self.tab("Doc2Txt"), text="Browse", command=self.browse_file_d2t)
        self.d2t_browse_button.grid(row=2, column=1, padx=10, pady=(20, 0))
        self.d2tsubmit_button = customtkinter.CTkButton(self.tab("Doc2Txt"), text="Convert", command=self.convert_docx_to_txt)
        self.d2t_label.grid(row=0, column=0)
        self.d2tsubmit_button.grid(row=3, column=0, pady=(20,0))
        self.d2t_entry.grid(row=2, column=0, padx=(20, 0), pady=(20, 0),  sticky="nsew")

        self.progress_bar_3 = customtkinter.CTkProgressBar(self.tab("Doc2Txt"))
        self.progress_bar_3.grid(row=4, column=0, pady=(20, 0), sticky="ew")
        self.progress_bar_3.set(0)
        self.progress_bar_3.configure(mode="determinate")

    def browse_file_d2p(self):
        filetypes = [('Word Documents', '*.docx')]
        filename = filedialog.askopenfilename(title='Select a .docx file', 
                                              initialdir=os.getcwd(),
                                              filetypes=filetypes)
        self.d2p_entry.delete(0, tk.END)
        self.d2p_entry.insert(0, filename)

    def convert_docx_to_pdf(self):
        docx_path = self.d2p_entry.get()
        pdf_path = os.path.splitext(docx_path)[0]

        self.progressbar_1.start()
        self.progressbar_1.update_idletasks()  # Force update before subprocess call

        # Schedule the subprocess call after a brief delay
        self.after(1500, self.execute_conversion, docx_path, pdf_path)


    def execute_conversion(self, docx_path, pdf_path):
        try:
            subprocess.call([r"C:\Program Files\LibreOffice\program\soffice.exe",
                '--convert-to',
                'pdf',
                '--outdir',
                pdf_path,
                docx_path])

        finally:
            # Stop the progress bar regardless of the conversion result
            self.progressbar_1.stop()
            self.progressbar_1.set(0)

    def browse_file_p2d(self):
        filetypes = [('Portable Document Format', '*.pdf')]
        filename = filedialog.askopenfilename(title='Select a .pdf file', 
                                              initialdir=os.getcwd(),
                                              filetypes=filetypes)
        self.p2d_entry.delete(0, tk.END)
        self.p2d_entry.insert(0, filename)


    def convert_pdf_to_docx(self):
        pdf_path = self.p2d_entry.get()
        docx_path = os.path.splitext(pdf_path)[0] + ".docx"

        self.progressbar_2.start()
        self.progressbar_2.update_idletasks()  # Force update before subprocess call

        # Schedule the subprocess call after a brief delay
        self.after(1500, self.execute_conversion_p2d, pdf_path, docx_path)

    def execute_conversion_p2d(self, pdf_path, docx_path):
        """
    Converts a PDF file to a DOCX file.

    Args:
        pdf_path (str): Path to the PDF file to convert.
        docx_path (str): Path to the resulting DOCX file.
    """
        try:

            # Open the PDF file using PyPDF2
            with open(pdf_path, 'rb') as pdf_file:
                pdf_reader = PdfReader(pdf_file)

                # Create a new DOCX file using the docx library
                docx_file = docx.Document()

                # Iterate through the pages of the PDF file
                for page in pdf_reader.pages:
                    # Get the text from the page
                    text = page.extract_text()

                    # had to remove File from PdfReader and extractText -> t_t

                    # Add the text to the DOCX file as a paragraph
                    docx_file.add_paragraph(text)

                # Save the DOCX file to disk
                docx_file.save(docx_path)

        finally:
            # Stop the progress bar regardless of the conversion result
            self.progressbar_2.stop()
            self.progressbar_2.set(0)

    def browse_file_d2t(self):
        filetypes = [('Word Documents', '*.docx')]
        filename = filedialog.askopenfilename(title='Select a .docx file', 
                                            initialdir=os.getcwd(),
                                            filetypes=filetypes)
        self.d2t_entry.delete(0, tk.END)
        self.d2t_entry.insert(0, filename)

    def convert_docx_to_txt(self):
        docx_path = self.d2t_entry.get()
        txt_path = os.path.splitext(docx_path)[0] + ".txt"

        # Use the `python-docx` library to convert DOCX to TXT
        doc = docx.Document(docx_path)
        text = []
        for para in doc.paragraphs:
            text.append(para.text)
        with open(txt_path, 'w') as f:
            f.write('\n'.join(text))

        self.progress_bar_3.start()
        self.progress_bar_3.update_idletasks()  # Force update before subprocess call

        # Schedule the subprocess call after a brief delay
        self.after(1500, self.execute_conversion_d2t, docx_path, txt_path)

    def execute_conversion_d2t(self, docx_path, txt_path):
        try:
            # No subprocess call needed for DOCX to TXT conversion
            pass

        finally:
            # Stop the progress bar regardless of the conversion result
            self.progress_bar_3.stop()